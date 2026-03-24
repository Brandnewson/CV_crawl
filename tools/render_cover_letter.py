"""Render a cover letter to DOCX using python-docx.

Input JSON schema:
{
    "name":           "Branson Tay",
    "address_line1":  "14 Samara West Mount",
    "address_line2":  "59-61 Clarendon Road",
    "address_line3":  "Leeds, West Yorkshire LS2 9NZ",
    "email":          "bransontay@gmail.com",
    "company_name":   "McLaren Racing",
    "company_address": "Woking, United Kingdom",   # optional, may be empty
    "date":           "14th March 2026",
    "salutation":     "Dear Hiring Manager,",
    "intro":          "...",
    "para1":          "...",
    "para2":          "...",
    "para3":          "...",
    "conclusion":     "...",
    "company":        "McLaren_Racing",            # safe filename slug
    "role":           "Junior_Data_Engineer"       # safe filename slug (max 30 chars)
}

Usage:
    uv run --project "C:/Code/CV_crawl" \\
        python "C:/Code/CV_crawl/tools/render_cover_letter.py" <json_path>

Writes DOCX to:
    C:\\Users\\brans\\OneDrive - University of Leeds\\GraduateJobHunting\\claude-cv-outputs\\
    {Company}_{Role}_CoverLetter_{YYYYMMDD}.docx

Prints the output path to stdout.
"""

import json
import sys
from datetime import datetime
from pathlib import Path

OUTPUT_DIR = Path(
    r"C:\Users\brans\OneDrive - University of Leeds\GraduateJobHunting\claude-cv-outputs"
)

LAYOUT_PRESETS = {
    "default": {
        "margin_cm": 2.5,
        "line_spacing": 1.15,
        "header_space_after": 2,
        "body_space_after": 8,
        "signoff_gap": 20,
    },
    "tight": {
        "margin_cm": 2.2,
        "line_spacing": 1.12,
        "header_space_after": 1,
        "body_space_after": 6,
        "signoff_gap": 16,
    },
    "tighter": {
        "margin_cm": 2.0,
        "line_spacing": 1.10,
        "header_space_after": 1,
        "body_space_after": 5,
        "signoff_gap": 14,
    },
}


def _set_font(run, name: str, size_pt: float, bold: bool = False) -> None:
    from docx.shared import Pt
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def render(data: dict) -> Path:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    preset_name = data.get("layout_preset", "default")
    preset = LAYOUT_PRESETS.get(preset_name, LAYOUT_PRESETS["default"])

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(preset["margin_cm"])
        section.bottom_margin = Cm(preset["margin_cm"])
        section.left_margin = Cm(preset["margin_cm"])
        section.right_margin = Cm(preset["margin_cm"])

    # ── Helper: styled paragraph ─────────────────────────────────────────────
    def para(text: str, size: float = 11, bold: bool = False,
             space_before: float = 0, space_after: float = 5,
             align: str = "left") -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = preset["line_spacing"]

        if align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        r = p.add_run(text)
        r.font.name = "Garamond"
        r.font.size = Pt(size)
        r.font.bold = bold

    # ── Sender details (name left, contact lines right) ───────────────────────
    para(
        data.get("name", "Branson Tay"),
        size=12,
        bold=True,
        space_after=preset["header_space_after"],
        align="left",
    )

    for field in ("address_line1", "address_line2", "address_line3"):
        val = data.get(field, "")
        if val:
            para(val, space_after=preset["header_space_after"], align="right")

    email = data.get("email", "")
    if email:
        para(email, space_after=preset["body_space_after"], align="right")

    # ── Company block onward (left aligned) ──────────────────────────────────
    company_name = data.get("company_name", "")
    if company_name:
        para(company_name, space_after=preset["header_space_after"], align="left")

    company_address = data.get("company_address", "")
    if company_address:
        para(company_address, space_after=preset["header_space_after"], align="left")

    # ── Date ─────────────────────────────────────────────────────────────────
    para(
        data.get("date", datetime.today().strftime("%d %B %Y").lstrip("0")),
        space_after=10,
        align="left",
    )

    # ── Salutation ───────────────────────────────────────────────────────────
    para(data.get("salutation", "Dear Hiring Manager,"), space_after=8, align="left")

    # ── Body paragraphs ──────────────────────────────────────────────────────
    for key in ("intro", "para1", "para2", "para3", "conclusion"):
        text = data.get(key, "")
        if text:
            para(text, space_before=0, space_after=preset["body_space_after"], align="left")

    # ── Sign-off ─────────────────────────────────────────────────────────────
    para("Yours Faithfully,", space_before=4, space_after=preset["signoff_gap"], align="left")
    para(data.get("name", "Branson Tay"), size=11, bold=False, space_after=0, align="left")

    # ── Output path ──────────────────────────────────────────────────────────
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_company = data.get("company", "Company")[:40]
    safe_role    = data.get("role", "Role")[:30]
    datestamp    = datetime.now().strftime("%Y%m%d")
    out_path     = OUTPUT_DIR / f"{safe_company}_{safe_role}_CoverLetter_{datestamp}.docx"

    doc.save(str(out_path))
    return out_path


def main() -> None:
    if len(sys.argv) < 2:
        print(f"Usage: python {sys.argv[0]} <json_path>", file=sys.stderr)
        sys.exit(1)

    json_path = Path(sys.argv[1])
    if not json_path.exists():
        print(f"ERROR: JSON file not found: {json_path}", file=sys.stderr)
        sys.exit(1)

    data = json.loads(json_path.read_text(encoding="utf-8"))
    out_path = render(data)
    print(str(out_path))


if __name__ == "__main__":
    main()
